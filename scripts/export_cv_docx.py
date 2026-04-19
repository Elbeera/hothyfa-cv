from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "content" / "cv-source.md"
OUTPUT = ROOT / "public" / "cv.docx"


def clean_inline_markdown(text: str) -> str:
    return text.replace("**", "").replace("`", "").strip()


def build_docx() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Missing source file: {SOURCE}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    for raw in lines:
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(clean_inline_markdown(line[4:]), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(clean_inline_markdown(line[3:]), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(clean_inline_markdown(line[2:]), level=1)
            continue
        if line.startswith("- "):
            text = clean_inline_markdown(line[2:])
            doc.add_paragraph(text, style="List Bullet")
            continue

        paragraph = doc.add_paragraph(clean_inline_markdown(line))
        paragraph.paragraph_format.space_after = Pt(6)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT}")


if __name__ == "__main__":
    build_docx()
